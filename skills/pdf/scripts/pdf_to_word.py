#!/usr/bin/env python3
"""
PDF to Word Converter - High Fidelity Version
Convert PDF to DOCX with maximum layout and format preservation.

Uses multiple backends for best results:
1. pdf2docx - Direct PDF to DOCX conversion (good for simple layouts)
2. Docling + Markdown-to-DOCX - Best for academic papers (highest accuracy)
3. pdfplumber fallback - Simple text extraction

Based on 2026 research: Docling (57.8k stars) is the most accurate PDF parser.
"""

import argparse
import os
import re
import sys
import time
from pathlib import Path
from typing import List, Optional, Tuple

# Check available backends
PDF2DOCX_AVAILABLE = False
DOCLING_AVAILABLE = False
PDFPLUMBER_AVAILABLE = False

try:
    from pdf2docx import Converter
    PDF2DOCX_AVAILABLE = True
except ImportError:
    pass

try:
    from docling.document_converter import DocumentConverter as DoclingConverter
    DOCLING_AVAILABLE = True
except ImportError:
    pass

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    pass

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    print("Error: python-docx not installed.")
    print("Please install: pip install python-docx")
    sys.exit(1)


def convert_with_docling(
    pdf_path: str,
    output_path: str,
    pages: Optional[List[int]] = None
) -> Tuple[bool, List[str], float]:
    """
    Convert PDF using Docling (highest accuracy for academic papers).

    Docling uses deep learning models for:
    - Layout analysis
    - Table structure recognition
    - Reading order detection
    - OCR (if needed)

    Returns Markdown which is then converted to DOCX.
    """
    errors = []

    if not DOCLING_AVAILABLE:
        errors.append("Docling not available. Install with: pip install docling")
        return False, errors, 0

    try:
        print("[Docling] Starting high-accuracy PDF parsing...")
        print("[Docling] Using deep learning models for layout analysis...")

        start_time = time.time()

        # Convert PDF using Docling
        converter = DoclingConverter()
        result = converter.convert(pdf_path)

        # Save as Markdown first
        temp_md = output_path.replace('.docx', '_temp.md')
        result.document.save_as_markdown(temp_md)

        elapsed = time.time() - start_time
        print(f"[Docling] PDF parsing completed in {elapsed:.2f}s")

        # Convert Markdown to DOCX
        print("[Docling] Converting Markdown to DOCX...")
        md_to_docx(temp_md, output_path)

        # Clean up temp file
        if os.path.exists(temp_md):
            os.remove(temp_md)

        print(f"[Docling] DOCX saved to: {output_path}")
        return True, errors, elapsed

    except Exception as e:
        errors.append(f"Docling conversion failed: {str(e)}")
        return False, errors, 0


def md_to_docx(md_path: str, docx_path: str):
    """
    Convert Markdown file to DOCX with proper formatting.
    """
    with open(md_path, 'r', encoding='utf-8') as f:
        md_content = f.read()

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)

    lines = md_content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i]

        # Skip empty lines
        if not line.strip():
            i += 1
            continue

        # Handle images (skip or add placeholder)
        if line.strip() == '<!-- image -->':
            i += 1
            continue

        # Handle headers
        if line.startswith('## '):
            text = line[3:].strip()
            h = doc.add_heading(text, level=2)
            i += 1
            continue

        if line.startswith('### '):
            text = line[4:].strip()
            h = doc.add_heading(text, level=3)
            i += 1
            continue

        if line.startswith('# '):
            text = line[2:].strip()
            h = doc.add_heading(text, level=1)
            i += 1
            continue

        # Handle lists
        if line.strip().startswith('- '):
            text = line.strip()[2:]
            doc.add_paragraph(text, style='List Bullet')
            i += 1
            continue

        # Handle numbered lists
        list_match = re.match(r'^(\d+)\.\s+(.+)$', line.strip())
        if list_match:
            text = list_match.group(2)
            doc.add_paragraph(text, style='List Number')
            i += 1
            continue

        # Handle links
        link_match = re.match(r'\[([^\]]+)\]\(([^)]+)\)', line.strip())
        if link_match:
            text = link_match.group(1)
            url = link_match.group(2)
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.underline = True
            run.font.color.rgb = RGBColor(0, 0, 255)
            i += 1
            continue

        # Handle tables (Markdown format)
        if '|' in line and i + 1 < len(lines) and '---' in lines[i + 1]:
            # Parse table
            rows = []
            while i < len(lines) and '|' in lines[i]:
                row = [cell.strip() for cell in lines[i].split('|')[1:-1]]
                if row and not all(c == '-' * len(c) for c in row):
                    rows.append(row)
                i += 1

            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = 'Table Grid'
                for row_idx, row_data in enumerate(rows):
                    for col_idx, cell_data in enumerate(row_data):
                        table.cell(row_idx, col_idx).text = cell_data
            continue

        # Regular paragraph
        doc.add_paragraph(line.strip())
        i += 1

    doc.save(docx_path)


def convert_with_pdf2docx(
    pdf_path: str,
    output_path: str,
    start_page: int = 0,
    end_page: int = None
) -> Tuple[bool, List[str], float]:
    """
    Convert PDF using pdf2docx (direct conversion, good for simple layouts).
    """
    errors = []

    if not PDF2DOCX_AVAILABLE:
        errors.append("pdf2docx not available. Install with: pip install pdf2docx")
        return False, errors, 0

    try:
        print("[pdf2docx] Starting direct PDF to DOCX conversion...")

        # Get total pages using pdfplumber
        if PDFPLUMBER_AVAILABLE:
            with pdfplumber.open(pdf_path) as pdf:
                total_pages = len(pdf.pages)
        else:
            total_pages = end_page if end_page else 100

        if end_page is None:
            end_page = total_pages

        start_time = time.time()

        cv = Converter(pdf_path)
        cv.convert(output_path, start=start_page, end=end_page)
        cv.close()

        elapsed = time.time() - start_time
        print(f"[pdf2docx] Conversion completed in {elapsed:.2f}s")
        print(f"[pdf2docx] DOCX saved to: {output_path}")

        return True, errors, elapsed

    except Exception as e:
        errors.append(f"pdf2docx conversion failed: {str(e)}")
        return False, errors, 0


def convert_with_pdfplumber(
    pdf_path: str,
    output_path: str,
    pages: Optional[List[int]] = None
) -> Tuple[int, List[str], float]:
    """
    Fallback: Convert PDF using pdfplumber + python-docx.
    Simple text extraction without advanced layout preservation.
    """
    errors = []
    pages_processed = 0

    if not PDFPLUMBER_AVAILABLE:
        errors.append("pdfplumber not available. Install with: pip install pdfplumber")
        return 0, errors, 0

    try:
        print("[pdfplumber] Starting simple text extraction...")

        start_time = time.time()
        doc = Document()

        pdf_name = Path(pdf_path).stem
        doc.add_heading(pdf_name, 0)

        with pdfplumber.open(pdf_path) as pdf:
            total_pages = len(pdf.pages)

            if pages is None:
                pages_to_process = range(total_pages)
            else:
                pages_to_process = [p for p in pages if 0 <= p < total_pages]

            for page_num in pages_to_process:
                page = pdf.pages[page_num]

                if page_num > 0:
                    doc.add_page_break()

                # Page marker
                marker = doc.add_paragraph()
                run = marker.add_run(f"--- Page {page_num + 1} ---")
                run.italic = True
                run.font.color.rgb = RGBColor(128, 128, 128)

                # Extract text
                text = page.extract_text()
                if text:
                    for line in text.split('\n'):
                        if line.strip():
                            doc.add_paragraph(line.strip())

                # Extract tables
                try:
                    tables = page.extract_tables()
                    for table in tables:
                        if table and len(table) > 0 and len(table[0]) > 0:
                            doc_table = doc.add_table(rows=len(table), cols=len(table[0]))
                            doc_table.style = 'Table Grid'
                            for row_idx, row in enumerate(table):
                                for col_idx, cell in enumerate(row):
                                    doc_table.cell(row_idx, col_idx).text = str(cell) if cell else ""
                except:
                    pass

                pages_processed += 1

        doc.save(output_path)
        elapsed = time.time() - start_time

        print(f"[pdfplumber] Conversion completed in {elapsed:.2f}s")
        print(f"[pdfplumber] DOCX saved to: {output_path}")

        return pages_processed, errors, elapsed

    except Exception as e:
        errors.append(f"pdfplumber conversion failed: {str(e)}")
        return pages_processed, errors, 0


def parse_page_range(page_range: str, total_pages: int) -> Tuple[int, int]:
    """Parse page range string into start and end page indices."""
    pages = []
    parts = page_range.split(',')

    for part in parts:
        part = part.strip()
        if '-' in part:
            start, end = part.split('-', 1)
            start = int(start.strip())
            end = int(end.strip())
            for p in range(start - 1, min(end, total_pages)):
                if p >= 0:
                    pages.append(p)
        else:
            p = int(part) - 1
            if 0 <= p < total_pages:
                pages.append(p)

    if not pages:
        return 0, total_pages

    return min(pages), max(pages) + 1


def main():
    parser = argparse.ArgumentParser(
        description="Convert PDF to Word with maximum fidelity.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  %(prog)s document.pdf -o document.docx
  %(prog)s document.pdf -o document.docx --backend docling
  %(prog)s document.pdf -o document.docx --backend pdf2docx
  %(prog)s document.pdf -o document.docx --pages 1-10

Backends:
  docling (default): Highest accuracy, uses deep learning models.
                     Best for academic papers, complex layouts.
                     Requires: pip install docling

  pdf2docx: Direct conversion, good for simple layouts.
            Requires: pip install pdf2docx

  pdfplumber: Simple text extraction, fallback option.
              Requires: pip install pdfplumber

Research (2026):
  - Docling: 57.8k stars, MIT license, IBM Research
  - pdf2docx: 3.4k stars, no longer maintained
  - pdfplumber: 10.1k stars, active maintenance
"""
    )

    parser.add_argument("pdf_file", help="Path to the PDF file")
    parser.add_argument("-o", "--output", dest="output_file", required=True,
                        help="Output Word document path")
    parser.add_argument("-p", "--pages", dest="page_range",
                        help="Page range (e.g., '1-5', '1,3,5')")
    parser.add_argument("--backend", choices=["docling", "pdf2docx", "pdfplumber", "auto"],
                        default="auto", help="Conversion backend (default: auto)")

    args = parser.parse_args()

    # Validate input
    if not os.path.exists(args.pdf_file):
        print(f"Error: File not found: {args.pdf_file}")
        sys.exit(1)

    # Get page count
    if PDFPLUMBER_AVAILABLE:
        with pdfplumber.open(args.pdf_file) as pdf:
            total_pages = len(pdf.pages)
    else:
        total_pages = 0
        print("Warning: Cannot determine page count without pdfplumber")

    print(f"\n{'='*60}")
    print(f"PDF to Word Converter (High Fidelity)")
    print(f"{'='*60}")
    print(f"Input: {args.pdf_file}")
    print(f"Output: {args.output_file}")
    print(f"Total pages: {total_pages}")

    # Determine backend
    backend = args.backend
    if backend == "auto":
        if DOCLING_AVAILABLE:
            backend = "docling"
            print(f"Backend: docling (auto-selected, highest accuracy)")
        elif PDF2DOCX_AVAILABLE:
            backend = "pdf2docx"
            print(f"Backend: pdf2docx (auto-selected)")
        else:
            backend = "pdfplumber"
            print(f"Backend: pdfplumber (auto-selected, fallback)")
    else:
        print(f"Backend: {backend} (user-specified)")

    print(f"{'='*60}\n")

    # Determine page range
    start_page = 0
    end_page = total_pages if total_pages > 0 else None

    if args.page_range and total_pages > 0:
        start_page, end_page = parse_page_range(args.page_range, total_pages)
        print(f"Converting pages {start_page + 1} to {end_page}")

    # Convert
    success = False
    errors = []
    elapsed = 0

    if backend == "docling":
        success, errors, elapsed = convert_with_docling(
            args.pdf_file, args.output_file
        )
    elif backend == "pdf2docx":
        success, errors, elapsed = convert_with_pdf2docx(
            args.pdf_file, args.output_file, start_page, end_page
        )
    else:  # pdfplumber
        pages_list = None if not args.page_range else list(range(start_page, end_page))
        _, errors, elapsed = convert_with_pdfplumber(
            args.pdf_file, args.output_file, pages_list
        )
        success = os.path.exists(args.output_file)

    # Verify output
    print(f"\n{'='*60}")
    if os.path.exists(args.output_file):
        file_size = os.path.getsize(args.output_file)
        print(f"SUCCESS: Output file created")
        print(f"File: {args.output_file}")
        print(f"Size: {file_size / 1024:.2f} KB")
        print(f"Time: {elapsed:.2f} seconds")
    else:
        print(f"FAILED: Output file not created")
        sys.exit(1)

    if errors:
        print(f"\nWarnings/Errors:")
        for error in errors:
            print(f"  - {error}")

    print(f"{'='*60}\n")


if __name__ == "__main__":
    main()
